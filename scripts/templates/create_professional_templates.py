"""
Create Professional Legal Document Templates
- Add header with logo
- Professional formatting with margins
- Legal document styling (stamp paper look)
- Proper colors and spacing
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_border_to_section(section):
    """Add decorative border to make it look like stamp paper"""
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn('w:pgBorders'))
    if pgBorders is None:
        pgBorders = OxmlElement('w:pgBorders')
        sectPr.append(pgBorders)
        pgBorders.set(qn('w:offsetFrom'), 'page')

    # Add borders (top, bottom, left, right)
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
        border = OxmlElement(border_name)
        border.set(qn('w:val'), 'double')  # Double line for professional look
        border.set(qn('w:sz'), '24')  # Size
        border.set(qn('w:space'), '24')  # Space from text
        border.set(qn('w:color'), '8B4513')  # Brown color for legal look
        pgBorders.append(border)

def create_professional_template(template_name, content_structure):
    """Create a professionally formatted legal template"""

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.3)
        section.right_margin = Inches(1.3)
        # Add stamp paper border
        add_border_to_section(section)

    # ===== HEADER WITH LOGO =====
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]

    # Try to add logo (if exists)
    logo_path = "./Lawmate/Lawmate/public/placeholder-logo.png"
    if os.path.exists(logo_path):
        run = header_para.add_run()
        run.add_picture(logo_path, width=Inches(1.5))
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add app name in header
    header_name = header.add_paragraph()
    header_run = header_name.add_run("Lawmate - Legal Document")
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(139, 69, 19)  # Brown color
    header_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add horizontal line
    header_line = header.add_paragraph()
    header_line_run = header_line.add_run("_" * 80)
    header_line_run.font.color.rgb = RGBColor(139, 69, 19)
    header_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== MAIN TITLE =====
    title = doc.add_paragraph()
    title_run = title.add_run(content_structure['title'])
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.underline = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_before = Pt(12)
    title.space_after = Pt(12)

    # Add spacing
    doc.add_paragraph()

    # ===== SUBTITLE =====
    if 'subtitle' in content_structure:
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(content_structure['subtitle'])
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.space_after = Pt(6)

    # ===== APPLICATION TYPE =====
    if 'application_type' in content_structure:
        app_type = doc.add_paragraph()
        app_type_run = app_type.add_run(content_structure['application_type'])
        app_type_run.font.size = Pt(12)
        app_type_run.font.bold = True
        app_type_run.font.color.rgb = RGBColor(139, 69, 19)
        app_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        app_type.space_after = Pt(12)

    doc.add_paragraph()  # Spacing

    # ===== HONORIFIC =====
    if 'honorific' in content_structure:
        honor = doc.add_paragraph()
        honor_run = honor.add_run(content_structure['honorific'])
        honor_run.font.size = Pt(11)
        honor.paragraph_format.line_spacing = 1.5
        honor.space_after = Pt(12)

    # ===== MAIN CONTENT SECTIONS =====
    for section_content in content_structure.get('sections', []):
        para = doc.add_paragraph()
        para_run = para.add_run(section_content)
        para_run.font.size = Pt(11)
        para_run.font.name = 'Times New Roman'
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Inches(0.3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ===== PRAYER SECTION =====
    if 'prayer' in content_structure:
        doc.add_paragraph()  # Spacing
        prayer_heading = doc.add_paragraph()
        prayer_heading_run = prayer_heading.add_run("PRAYER:")
        prayer_heading_run.font.size = Pt(12)
        prayer_heading_run.font.bold = True
        prayer_heading.space_before = Pt(12)
        prayer_heading.space_after = Pt(6)

        for prayer_item in content_structure['prayer']:
            prayer_para = doc.add_paragraph()
            prayer_para_run = prayer_para.add_run(prayer_item)
            prayer_para_run.font.size = Pt(11)
            prayer_para.paragraph_format.line_spacing = 1.5
            prayer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ===== FOOTER SECTION (Applicant/Counsel) =====
    if 'footer' in content_structure:
        doc.add_paragraph()
        doc.add_paragraph()

        for footer_line in content_structure['footer']:
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run(footer_line)
            footer_run.font.size = Pt(11)
            footer_run.font.bold = True if 'APPLICANT' in footer_line or 'Through Counsel' in footer_line else False
            footer_para.space_after = Pt(3)

    # ===== PAGE FOOTER =====
    footer_section = doc.sections[0].footer
    footer_para = footer_section.paragraphs[0]
    footer_para.text = "Generated by Lawmate | www.lawmate.com"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    return doc

# Create Consolidation of Cases Template
consolidation_structure = {
    'title': 'IN THE COURT OF {COURT_NAME}, {CITY_DISTRICT}.',
    'subtitle': 'Case Title: {CASE_TITLE}, e.g., {APPLICANT_NAME} versus {RESPONDENT_NAME}',
    'application_type': 'APPLICATION FOR CONSOLIDATION OF CASES',
    'honorific': 'May it please Your Honour,',
    'sections': [
        'That the Applicant/Petitioner is a party in multiple legal proceedings concerning the same subject matter, specifically {SUBJECT_MATTER}.',
        '',
        'That the Opposite Party/Respondent(s) have filed different cases against the Applicant/Petitioner in different courts, as detailed below:',
        '',
        'A. {CASE_TITLE_1}, pending in the Court of {COURT_NAME_1}, {LOCATION_1}. The next date of hearing is fixed for {HEARING_DATE_1}.',
        '',
        'B. {CASE_TITLE_2}, pending in the Court of {COURT_NAME_2}, {LOCATION_2}. The next date of hearing is fixed for {HEARING_DATE_2}.',
        '',
        'C. {CASE_TITLE_3}, pending in the Court of {COURT_NAME_3}, {LOCATION_3}. The next date of hearing is fixed for {HEARING_DATE_3}.',
        '',
        'That the core issues, parties, and subject matter involved in the aforementioned cases are identical/substantially the same.',
        '',
        'That the Respondent(s) have initiated these multiple proceedings in different courts with a mala fide intention and ulterior motive, solely to harass the Applicant and to prolong the litigation, as the cases are frivolous and baseless.',
        '',
        'That if the above-mentioned cases are not consolidated and heard together by the same court, the Applicant will suffer irreparable loss and injury, and it will lead to conflicting judgments and a waste of precious judicial time.',
    ],
    'prayer': [
        'Under the circumstances stated above, it is most humbly prayed that this Honourable Court may graciously be pleased to:',
        '',
        'Direct the consolidation of the following cases:',
        '(i) {CASE_TITLE_1}',
        '(ii) {CASE_TITLE_2}',
        '(iii) {CASE_TITLE_3}',
        '',
        'and order that they be heard and decided together by this Honourable Court / by a single court of competent jurisdiction, in the interest of justice, equity, and to avoid multiplicity of proceedings.',
    ],
    'footer': [
        'APPLICANT / PETITIONER',
        '',
        'Through Counsel:',
        '{LAWYER_NAME}',
        'Advocate',
        'Enrollment No. {ENROLLMENT_NUMBER}',
        'Dated: {DOCUMENT_DATE}',
    ]
}

print("Creating professional Consolidation of Cases template...")
doc = create_professional_template("Consolidation of cases", consolidation_structure)
output_path = "./data/DOC_TEMPLATE/Consolidation of cases.docx"
doc.save(output_path)
print(f"✅ Created: {output_path}")

print("\n✅ Professional templates created successfully!")
print("All templates now have:")
print("  - Logo header")
print("  - Stamp paper borders (brown double lines)")
print("  - Proper margins (1.2\" - 1.3\")")
print("  - Professional fonts and spacing")
print("  - Legal document styling")
print("  - Color-coded sections")
